import random
import os
import string
import datetime
from typing import Literal
from docx import Document

from src.utils.document_style import DocumentStyleManager


certificate_types = Literal["course", "category"]
doc_manager = DocumentStyleManager()


class CertificateWriter:
    def __init__(
            self,
            cert_type: certificate_types,
            student_name: str,
            category_name: str,
            course_name: str = None,
            courses_list: list = None,
    ):
        self.cert_type = cert_type
        self.student_name = student_name
        self.course_name = course_name
        self.category_name = category_name
        self.courses_list = courses_list

    def generate_certificate_seria(self):
        seria_num: list = random.choices(string.digits, k=4)

        if self.cert_type == "course":
            return f"M01 OC-06/{''.join(seria_num)}"

        return f"M01 OP-06/{''.join(seria_num)}"

    @staticmethod
    def get_actual_date():
        date = datetime.date.today()
        return date.strftime("%d.%m.%Y")


    def write_course_certificate_data(self):
        doc = Document('templates/course_cert.docx')
        seria = self.generate_certificate_seria()

        for paragraph in doc.paragraphs:
            if "{seria}" in paragraph.text:
                new_text = paragraph.text.replace("{seria}", seria)
                paragraph.text = new_text
                color = doc_manager.hex_to_rgb("#03183F")
                doc_manager.set_paragraph_text_color(paragraph, *color)
                doc_manager.set_paragraph_font_size(paragraph, 9)
                doc_manager.set_paragraph_bold(paragraph)
                doc_manager.set_paragraph_font_name(paragraph, "Inter")
                doc_manager.set_paragraph_alignment(paragraph, 'left')

            if "{date}" in paragraph.text:
                new_text = paragraph.text.replace("{date}", self.get_actual_date())
                paragraph.text = new_text
                color = doc_manager.hex_to_rgb("#03183F")
                doc_manager.set_paragraph_text_color(paragraph, *color)
                doc_manager.set_paragraph_font_size(paragraph, 8)
                doc_manager.set_paragraph_bold(paragraph)
                doc_manager.set_paragraph_font_name(paragraph, "Inter")
                doc_manager.set_paragraph_alignment(paragraph, 'center')

            if "{student_name}" in paragraph.text:
                new_text = paragraph.text.replace("{student_name}", self.student_name)
                paragraph.text = new_text
                color = doc_manager.hex_to_rgb("#0054CA")
                doc_manager.set_paragraph_text_color(paragraph, *color)
                doc_manager.set_paragraph_font_size(paragraph, 36)
                doc_manager.set_paragraph_font_name(paragraph, "Carattere")
                doc_manager.set_paragraph_alignment(paragraph, 'center')
                doc_manager.set_paragraph_italic(paragraph)

            if "{course_name}" in paragraph.text:
                new_text = paragraph.text.replace("{course_name}", self.course_name)
                paragraph.text = new_text
                color = doc_manager.hex_to_rgb("#001C54")
                doc_manager.set_paragraph_text_color(paragraph, *color)
                doc_manager.set_paragraph_font_size(paragraph, 20)
                doc_manager.set_paragraph_font_name(paragraph, "Inter")
                doc_manager.set_paragraph_alignment(paragraph, 'center')
                doc_manager.set_paragraph_italic(paragraph)

            if "{category_name}" in paragraph.text:
                new_text = paragraph.text.replace("{category_name}", self.category_name)
                paragraph.text = new_text
                color = doc_manager.hex_to_rgb("#001C54")
                doc_manager.set_paragraph_text_color(paragraph, *color)
                doc_manager.set_paragraph_font_size(paragraph, 14)
                doc_manager.set_paragraph_font_name(paragraph, "Inter")
                doc_manager.set_paragraph_alignment(paragraph, 'center')
                doc_manager.set_paragraph_italic(paragraph)

        dir_name = 'static/certificates/courses/'
        os.makedirs(dir_name, exist_ok=True)

        file_id = seria.replace(' ', '-')
        file_id = file_id.replace('/', '-')

        filename = f"{dir_name}{file_id}.docx"
        doc.save(filename)
        return dir_name, filename

    def write_category_certificate_data(self):
        doc = Document('templates/category_cert.docx')
        seria = self.generate_certificate_seria()

        for paragraph in doc.paragraphs:

            if "{seria}" in paragraph.text:
                new_text = paragraph.text.replace("{seria}", seria)
                paragraph.text = new_text
                color = doc_manager.hex_to_rgb("#03183F")
                doc_manager.set_paragraph_text_color(paragraph, *color)
                doc_manager.set_paragraph_font_size(paragraph, 9)
                doc_manager.set_paragraph_bold(paragraph)
                doc_manager.set_paragraph_font_name(paragraph, "Inter")
                doc_manager.set_paragraph_alignment(paragraph, 'left')

            if "{date}" in paragraph.text:
                new_text = paragraph.text.replace("{date}", self.get_actual_date())
                paragraph.text = new_text
                color = doc_manager.hex_to_rgb("#03183F")
                doc_manager.set_paragraph_text_color(paragraph, *color)
                doc_manager.set_paragraph_font_size(paragraph, 8)
                doc_manager.set_paragraph_bold(paragraph)
                doc_manager.set_paragraph_font_name(paragraph, "Inter")
                doc_manager.set_paragraph_alignment(paragraph, 'center')

            if "{student_name}" in paragraph.text:
                new_text = paragraph.text.replace("{student_name}", self.student_name)
                paragraph.text = new_text
                color = doc_manager.hex_to_rgb("#B90000")
                doc_manager.set_paragraph_text_color(paragraph, *color)
                doc_manager.set_paragraph_font_size(paragraph, 36)
                doc_manager.set_paragraph_font_name(paragraph, "Carattere")
                doc_manager.set_paragraph_alignment(paragraph, 'center')
                doc_manager.set_paragraph_italic(paragraph)

            if "{category_name}" in paragraph.text:
                new_text = paragraph.text.replace("{category_name}", self.category_name)
                paragraph.text = new_text
                color = doc_manager.hex_to_rgb("#001C54")
                doc_manager.set_paragraph_text_color(paragraph, *color)
                doc_manager.set_paragraph_font_size(paragraph, 20)
                doc_manager.set_paragraph_font_name(paragraph, "Inter")
                doc_manager.set_paragraph_alignment(paragraph, 'center')
                doc_manager.set_paragraph_italic(paragraph)

            if "{course_list}" in paragraph.text:
                new_p = paragraph.insert_paragraph_before()

                for i in range(0, len(self.courses_list), 2):
                    # Получаем текущую пару курсов
                    pair = self.courses_list[i:i + 2]

                    # Формируем строку для текущей пары курсов
                    for j, course in enumerate(pair):
                        if j > 0:  # Добавляем пробелы перед вторым элементом пары
                            new_p.add_run("      ")  # 4 пробела
                        new_p.add_run(u"\u25aa ")  # Добавляем символ
                        new_p.add_run(f"{course}")  # Добавляем название курса

                    # Добавляем перенос строки после каждой пары
                    new_p.add_run("\n")

                color = doc_manager.hex_to_rgb("#001C54")
                doc_manager.set_paragraph_text_color(new_p, *color)
                doc_manager.set_paragraph_font_size(new_p, 12)
                doc_manager.set_paragraph_font_name(new_p, "Inter")
                doc_manager.set_paragraph_alignment(new_p, 'center')
                doc_manager.set_paragraph_italic(new_p)
                doc_manager.set_paragraph_indent(new_p, 140, 140)

                new_text = paragraph.text.replace("{course_list}", "")
                paragraph.text = new_text

        dir_name = 'static/certificates/categories/'
        os.makedirs(dir_name, exist_ok=True)

        file_id = seria.replace(' ', '-')
        file_id = file_id.replace('/', '-')

        filename = f"{dir_name}{file_id}.docx"
        doc.save(filename)
        return dir_name, filename
